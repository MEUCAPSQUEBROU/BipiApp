# -*- coding: utf-8 -*-
"""Gera a ERS do Bipi preenchida, mantendo a estrutura do template da II APP JAM.

Saída: Documentação/ERS_Bipi.docx (editável). Não toca no .docx original.
Rodar: python tool/gen_ers.py
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = "Documentação/ERS_Bipi.docx"
ACCENT = RGBColor(0x2E, 0x74, 0xB5)  # azul dos títulos (como no template)

doc = Document()


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_toc(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = ("Clique com o botão direito aqui e escolha "
                        "“Atualizar campo” para gerar o sumário.")
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, placeholder, end):
        run._r.append(el)


def force_update_fields():
    """Faz o Word atualizar os campos (sumário) ao abrir o arquivo."""
    el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true")
    doc.settings.element.append(el)


def styled_title(text, size=18, color=ACCENT, space_before=10, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def body(text):
    return doc.add_paragraph(text)


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def label_line(label, value):
    p = doc.add_paragraph()
    p.add_run(f"{label} ").bold = True
    p.add_run(value)
    return p


# ───────────────────────── Cabeçalho (todas as páginas) ─────────────────────
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.text = ""
r = hp.add_run("UNINASSAU")
r.bold = True
hp.add_run("\t\tII APP JAM").bold = True
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ───────────────────────────────── Capa ─────────────────────────────────────
t = doc.add_paragraph()
tr = t.add_run("Especificação de Requisitos de Software (ERS)")
tr.bold = True
tr.font.size = Pt(20)
doc.add_paragraph()
label_line("Curso:", "[Informe o nome]")
label_line("Disciplina:", "[Informe o nome]")
label_line("Professor:", "[Informe o nome]")
label_line("Nome do Aplicativo:", "Bipi")
label_line("Matrícula/Nome:", "[Informe o nome]")

# ──────────────────────────────── Sumário ───────────────────────────────────
styled_title("Sumário", size=18)
add_toc(doc.add_paragraph())
doc.add_page_break()

# ─────────────────────────────── 1. Introdução ──────────────────────────────
doc.add_heading("1. Introdução", level=1)
body("O Bipi é um aplicativo educacional gamificado sobre segurança no trânsito, "
     "criado em apoio à campanha Maio Amarelo. Este documento especifica os "
     "requisitos funcionais e não funcionais do aplicativo, servindo de "
     "referência para o desenvolvimento, os testes e a validação do sistema.")

doc.add_heading("1.1 Objetivo", level=2)
body("Este documento descreve os requisitos funcionais e não funcionais do "
     "aplicativo Bipi, com o objetivo de orientar o desenvolvimento, os testes "
     "e a validação do sistema, além de alinhar o entendimento entre a equipe e "
     "os interessados (professores e avaliadores da II APP JAM).")

doc.add_heading("1.2 Escopo", level=2)
body("O Bipi tem como finalidade ensinar regras e atitudes de trânsito de forma "
     "rápida e divertida — no estilo “Duolingo do trânsito” —, voltado a "
     "estudantes do ensino médio. O aplicativo permite que o usuário crie uma "
     "conta, percorra uma trilha diária de fases, responda a quizzes de múltipla "
     "escolha com feedback imediato (explicação, confete, som e vibração), "
     "acumule pontos e acompanhe sua posição em um ranking. O conteúdo apoia a "
     "campanha Maio Amarelo de conscientização sobre segurança viária.")

doc.add_heading("1.3 Definições, Acrônimos e Abreviações", level=2)
bullets([
    "ERS: Especificação de Requisitos de Software",
    "UI: Interface do Usuário (User Interface)",
    "API: Interface de Programação de Aplicações",
    "CTB: Código de Trânsito Brasileiro",
    "SDK: Kit de Desenvolvimento de Software (Software Development Kit)",
    "Trilha: sequência diária de fases que o usuário percorre",
    "Streak: contagem de acertos consecutivos",
])

doc.add_heading("1.4 Referências", level=2)
bullets([
    "IEEE 830 — Padrão de documentação de requisitos de software",
    "Documentação oficial do Flutter (flutter.dev) e do Firebase "
    "(firebase.google.com)",
    "Código de Trânsito Brasileiro (Lei nº 9.503/1997)",
    "Movimento Maio Amarelo (maioamarelo.com)",
    "Repositório do projeto: github.com/MEUCAPSQUEBROU/BipiApp",
])

# ───────────────────────────── 2. Descrição Geral ───────────────────────────
doc.add_heading("2. Descrição Geral", level=1)

doc.add_heading("2.1 Perspectiva do Produto", level=2)
body("O Bipi é um aplicativo móvel independente (standalone) integrado a "
     "serviços externos do Firebase para autenticação e ranking. Foi "
     "desenvolvido em Flutter, a partir de uma única base de código compatível "
     "com:")
bullets([
    "Android (plataforma principal de distribuição)",
    "iOS",
    "Web",
])

doc.add_heading("2.2 Funções do Produto", level=2)
body("O sistema deverá permitir:")
bullets([
    "Cadastro e login de usuários (e-mail/senha e conta Google)",
    "Recuperação de senha por e-mail",
    "Onboarding de boas-vindas com o mascote Bipi",
    "Trilha diária de fases, com desbloqueio progressivo e rampa de dificuldade",
    "Quiz de múltipla escolha com feedback imediato (explicação, confete, "
    "flash, som e vibração)",
    "Contagem de streak e pontuação por acerto",
    "Ranking de jogadores com pódio (top 3) e leaderboard",
    "Configuração de áudio (ativar/silenciar efeitos sonoros), com preferência "
    "salva",
])

doc.add_heading("2.3 Características dos Usuários", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Tipo de usuário"
hdr[1].text = "Descrição"
for c in hdr:
    set_cell_shading(c, "BDD7EE")
    c.paragraphs[0].runs[0].bold = True
rows = [
    ("Estudante (usuário comum)",
     "Cria conta, percorre a trilha, responde aos quizzes e acompanha sua "
     "posição no ranking. Público-alvo: ensino médio."),
    ("Administrador / Educador",
     "Curadoria do banco de perguntas e acompanhamento pedagógico (gestão de "
     "conteúdo — evolução planejada)."),
]
for tipo, desc in rows:
    cells = table.add_row().cells
    cells[0].text = tipo
    cells[1].text = desc

doc.add_heading("2.4 Restrições", level=2)
bullets([
    "Necessita conexão com a internet para autenticação e ranking",
    "Desenvolvido em Flutter (Dart SDK ^3.12)",
    "Requer versões mínimas do sistema operacional (Android 6.0 / API 23 ou "
    "superior)",
    "Depende dos serviços do Firebase (Authentication e Cloud Firestore)",
])

doc.add_heading("2.5 Suposições e Dependências", level=2)
bullets([
    "Disponibilidade dos serviços Firebase (Authentication e Firestore)",
    "Acesso à internet pelo dispositivo do usuário",
    "Conta Google válida para login social (opcional)",
    "Projeto Firebase configurado (google-services.json no Android)",
])

# ─────────────────────────── 3. Requisitos Funcionais ───────────────────────
doc.add_heading("3. Requisitos Funcionais", level=1)

doc.add_heading("RF01 — Cadastro de Usuário", level=2)
body("O sistema deve permitir que o usuário crie uma conta informando:")
bullets(["Nome", "E-mail", "Senha", "Confirmação de senha"])
body("Alternativamente, o cadastro pode ser feito com uma conta Google.")

doc.add_heading("RF02 — Login", level=2)
body("O sistema deve permitir autenticação com:")
bullets(["E-mail e senha", "Conta Google (login social)"])
body("Mensagens de erro são exibidas de forma amigável (e-mail inválido, senha "
     "incorreta, sem conexão etc.).")

doc.add_heading("RF03 — Recuperação de Senha", level=2)
body("O sistema deve permitir que o usuário redefina a senha, enviando um link "
     "de redefinição para o e-mail informado (via Firebase Authentication).")

doc.add_heading("RF04 — Onboarding", level=2)
body("Após o primeiro login, o sistema deve exibir uma introdução com o mascote "
     "Bipi, apresentando a proposta do app. A conclusão do onboarding é "
     "registrada localmente para não se repetir.")

doc.add_heading("RF05 — Trilha Diária", level=2)
body("O sistema deve apresentar uma trilha de fases que muda a cada dia. A "
     "conclusão de uma fase desbloqueia a próxima, e as fases seguem uma rampa "
     "de dificuldade. O progresso do dia é exibido (ex.: 3/5 fases).")

doc.add_heading("RF06 — Quiz", level=2)
body("O sistema deve apresentar perguntas de múltipla escolha. Para cada "
     "resposta, deve:")
bullets([
    "Indicar visualmente acerto/erro nas alternativas",
    "Exibir uma explicação curta sobre a resposta",
    "Tocar efeito sonoro e acionar vibração (háptico) conforme o resultado",
    "Atualizar o contador de streak (acertos consecutivos)",
    "Ao final, exibir a pontuação e o maior streak",
])
body("Cada acerto vale 10 pontos, somados ao ranking na primeira conclusão da "
     "fase.")

doc.add_heading("RF07 — Ranking", level=2)
body("O sistema deve exibir um ranking dos jogadores por pontuação "
     "(decrescente), com pódio (top 3, com foto e primeiro nome) e leaderboard. "
     "Os dados são sincronizados via Cloud Firestore (coleção “ranking”).")

doc.add_heading("RF08 — Efeitos Sonoros e Háptico", level=2)
body("O sistema deve fornecer feedback sonoro e tátil nos momentos-chave "
     "(acerto, erro, conclusão de fase, conclusão da trilha e toques de botão) e "
     "permitir silenciar os efeitos, salvando a preferência do usuário.")

# ─────────────────────────── 4. Requisitos Não Funcionais ───────────────────
doc.add_heading("4. Requisitos Não Funcionais", level=1)

doc.add_heading("4.1 Desempenho", level=2)
bullets([
    "A interface deve responder de forma fluida (alvo de 60 fps nas transições)",
    "Os efeitos sonoros são pré-carregados e reproduzidos em baixa latência",
    "As telas devem carregar em até 3 segundos em condições normais de rede",
])

doc.add_heading("4.2 Segurança", level=2)
bullets([
    "Autenticação gerenciada pelo Firebase Authentication",
    "Comunicação com os serviços via HTTPS",
    "Acesso aos dados controlado por regras de segurança do Firestore",
    "Senhas não são armazenadas pelo aplicativo (ficam a cargo do Firebase)",
])

doc.add_heading("4.3 Usabilidade", level=2)
bullets([
    "Linguagem e visual adequados ao ensino médio (sem infantilizar)",
    "Navegação simples e feedback imediato a cada ação",
    "Tempo de aprendizado reduzido; sessões curtas",
])

doc.add_heading("4.4 Confiabilidade", level=2)
bullets([
    "O ranking é best-effort: falhas de rede não interrompem o jogo",
    "Disponibilidade dependente do Firebase (SLA do provedor)",
])

doc.add_heading("4.5 Portabilidade", level=2)
bullets([
    "Base única em Flutter para Android, iOS e Web",
    "Compatível com Android 6.0 (API 23) ou superior",
])

# ─────────────────────────── 5. Requisitos de Interface ─────────────────────
doc.add_heading("5. Requisitos de Interface", level=1)

doc.add_heading("5.1 Interface do Usuário", level=2)
bullets([
    "Material Design 3, com a paleta da campanha Maio Amarelo",
    "Mascote Bipi com expressões (normal, feliz, dúvida, decepção)",
    "Layout responsivo e navegação por rotas (go_router)",
])

doc.add_heading("5.2 Interface com Hardware", level=2)
bullets([
    "Alto-falante: reprodução dos efeitos sonoros",
    "Motor de vibração: feedback háptico em acertos/erros",
    "Câmera e GPS: não utilizados na versão atual",
])

doc.add_heading("5.3 Interface com Software", level=2)
bullets([
    "Firebase Authentication (e-mail/senha e Google Sign-In)",
    "Cloud Firestore (ranking)",
    "shared_preferences (preferências locais: onboarding e mudo)",
    "Pacotes: audioplayers, confetti, google_fonts",
])

# ───────────────────────────────── 6. Modelagem ─────────────────────────────
doc.add_heading("6. Modelagem", level=1)

doc.add_heading("6.1 Casos de Uso", level=2)
body("Principais casos de uso:")
bullets([
    "UC01: Criar conta / Fazer login (e-mail ou Google)",
    "UC02: Visualizar onboarding",
    "UC03: Percorrer a trilha diária",
    "UC04: Responder a um quiz e receber feedback",
    "UC05: Consultar o ranking",
    "UC06: Ativar/silenciar os efeitos sonoros",
])

doc.add_heading("6.2 Diagramas", level=2)
body("Os diagramas a seguir devem ser anexados (ver seção 8):")
bullets([
    "Diagrama de casos de uso",
    "Diagrama de classes",
    "Diagrama de sequência (fluxo de login e de resposta a uma pergunta)",
])

# ───────────────────────── 7. Critérios de Aceitação ────────────────────────
doc.add_heading("7. Critérios de Aceitação", level=1)
body("Cada requisito deve incluir critérios claros. Exemplos:")

doc.add_paragraph().add_run("Exemplo (RF02 — Login):").bold = True
bullets([
    "Credenciais válidas → usuário é autenticado e direcionado à tela inicial",
    "E-mail/senha inválidos → mensagem de erro amigável é exibida",
])

doc.add_paragraph().add_run("Exemplo (RF06 — Quiz):").bold = True
bullets([
    "Resposta correta → alternativa fica verde, toca o som de acerto, dispara "
    "confete e incrementa o streak",
    "Resposta incorreta → alternativa fica vermelha, toca o som de erro, exibe "
    "a explicação e zera o streak",
])

doc.add_paragraph().add_run("Exemplo (RF07 — Ranking):").bold = True
bullets([
    "Concluir uma fase pela primeira vez → soma 10 pontos por acerto e atualiza "
    "a posição no ranking",
])

# ───────────────────────────── 8. Anexos (Opcional) ─────────────────────────
doc.add_heading("8. Anexos (Opcional)", level=1)
bullets([
    "Protótipos de tela (capturas do app: login, onboarding, trilha, quiz, "
    "ranking)",
    "Fluxos de navegação",
    "Repositório do projeto: github.com/MEUCAPSQUEBROU/BipiApp",
])

force_update_fields()
doc.core_properties.title = "Especificação de Requisitos de Software — Bipi"
doc.save(OUT)
print(f"Gerado: {OUT}")
